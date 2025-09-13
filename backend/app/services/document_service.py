from typing import List, Dict, Any, Optional
import os
import hashlib
import json
import asyncio
from datetime import datetime
from fastapi import UploadFile, HTTPException, status
from sqlalchemy.orm import Session
try:
    import pypdf
except ImportError:
    import PyPDF2 as pypdf
import docx
import re
from contextlib import asynccontextmanager

from app.models.document import Document
from app.services.vector_service import VectorService
from app.services.storage_service import StorageService
from app.core.tenant import validate_tenant_limits
from app.config.settings import settings

try:
    import google.generativeai as genai
except ImportError:
    genai = None

try:
    import openai
except ImportError:
    openai = None


class DocumentService:
    """
    Service for document management and processing
    """
    
    def __init__(self, db: Session):
        self.db = db
        self.vector_service = VectorService()
        self.storage_service = StorageService()
    
    async def upload_document(
        self,
        file: UploadFile,
        document_code: Optional[str],
        title: str,
        document_type: str,
        category: Optional[str],
        tenant_id: int,
        uploaded_by: int,
        industry_sectors: Optional[List[str]] = None,
        authority: Optional[str] = None,
        subcategory: Optional[str] = None,
        scope: Optional[str] = None,
        force_reload: bool = False
    ) -> Document:
        """
        Upload and process a document
        """
        # Validate file type
        file_extension = os.path.splitext(file.filename)[1].lower()
        allowed_extensions = ['.pdf', '.docx', '.doc', '.txt']
        
        if file_extension not in allowed_extensions:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"File type {file_extension} not allowed"
            )
        
        # Calculate file hash
        file_content = await file.read()
        file_hash = hashlib.sha256(file_content).hexdigest()
        await file.seek(0)  # Reset file pointer
        
        # Check for duplicate file by hash
        if not force_reload:
            existing_by_hash = self.db.query(Document).filter(
                Document.tenant_id == tenant_id,
                Document.file_hash == file_hash,
                Document.is_active == True
            ).first()
            
            if existing_by_hash:
                raise HTTPException(
                    status_code=status.HTTP_409_CONFLICT,
                    detail=f"File already exists as document '{existing_by_hash.title}' (Code: {existing_by_hash.document_code}, Version: {existing_by_hash.version}). Use force_reload=true to re-analyze."
                )
        
        # Check tenant limits
        from app.models.tenant import Tenant
        tenant = self.db.query(Tenant).filter(Tenant.id == tenant_id).first()
        if not validate_tenant_limits(tenant, "upload_document", self.db):
            raise HTTPException(
                status_code=status.HTTP_403_FORBIDDEN,
                detail="Document limit reached for tenant"
            )
        
        # Generate document_code if not provided
        if not document_code:
            # Generate code from title and timestamp
            base_code = re.sub(r'[^a-zA-Z0-9]', '_', title.lower())[:30]
            timestamp = datetime.now().strftime("%Y%m%d%H%M%S")
            document_code = f"{base_code}_{timestamp}"
        
        # Check if document already exists
        existing_doc = self.db.query(Document).filter(
            Document.tenant_id == tenant_id,
            Document.document_code == document_code
        ).first()
        
        if existing_doc:
            # Create new version
            version_parts = existing_doc.version.split('.')
            new_version = f"{version_parts[0]}.{int(version_parts[1]) + 1}"
        else:
            new_version = "1.0"
        
        try:
            # Upload file to storage
            file_path = await self.storage_service.upload_file(
                file, f"documents/{tenant_id}/{document_code}_{new_version}{file_extension}"
            )
            
            # Extract text content
            content = await self._extract_text_content(file, file_extension)
            
            # Extract metadata using AI if not provided
            ai_metadata = await self._extract_ai_metadata(content, title, document_type)
            
            # Use AI extracted values if not provided
            if not category:
                category = ai_metadata.get("category", "general")
            if not subcategory:
                subcategory = ai_metadata.get("subcategory")
            if not authority:
                authority = ai_metadata.get("authority")
            if not scope:
                scope = ai_metadata.get("scope", "tenant")
            if not industry_sectors:
                industry_sectors = ai_metadata.get("industry_sectors", [])
            
            # Extract keywords from full document (AI only, no fallback)
            keywords = ai_metadata.get("keywords", [])
            
            # If AI failed to extract keywords, fail the upload
            if not keywords:
                # Delete uploaded file
                await self.storage_service.delete_file(file_path)
                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail="Failed to extract keywords from document. Please ensure the document contains relevant HSE content or try again later."
                )
            
            # Create semantic chunks
            chunks = await self._create_semantic_chunks(
                content, document_type, document_code
            )
            
            # Create document record
            document = Document(
                tenant_id=tenant_id,
                document_code=document_code,
                title=title,
                document_type=document_type,
                category=category,
                subcategory=subcategory,
                scope=scope,
                industry_sectors=industry_sectors or [],
                keywords=keywords,
                file_path=file_path,
                file_hash=file_hash,
                content_summary=content[:500] + "..." if len(content) > 500 else content,
                version=new_version,
                authority=authority,
                uploaded_by=uploaded_by,
                publication_date=datetime.now().date()
            )
            
            self.db.add(document)
            self.db.commit()
            self.db.refresh(document)
            
            # Add chunks to vector database with timeout
            try:
                chunk_ids = await asyncio.wait_for(
                    self.vector_service.add_document_chunks(
                        document_id=document.id,
                        document_code=document_code,
                        title=title,
                        chunks=chunks,
                        document_type=document_type,
                        category=category,
                        tenant_id=tenant_id,
                        industry_sectors=industry_sectors,
                        authority=authority
                    ),
                    timeout=300.0  # 5 minutes timeout for vector processing
                )
            except asyncio.TimeoutError:
                # Clean up on timeout - delete document and file
                print(f"[DocumentService] Vector processing timed out for document {document.id}")
                
                # Delete document from database
                self.db.delete(document)
                self.db.commit()
                
                # Delete file from storage
                await self.storage_service.delete_file(file_path)
                
                raise HTTPException(
                    status_code=status.HTTP_408_REQUEST_TIMEOUT,
                    detail="Document processing timed out. The document is too large or complex. Document has been removed."
                )
            except Exception as vector_error:
                # Clean up on vector processing error
                print(f"[DocumentService] Vector processing failed for document {document.id}: {str(vector_error)}")
                
                # Delete document from database
                self.db.delete(document)
                self.db.commit()
                
                # Delete file from storage
                await self.storage_service.delete_file(file_path)
                
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"Failed to process document vectors: {str(vector_error)}. Document has been removed."
                )
            
            # Store first chunk ID as vector_id and calculate overall relevance score
            if chunk_ids:
                document.vector_id = chunk_ids[0]
                # Calculate overall document relevance score from chunks
                if chunks:
                    avg_relevance = sum(chunk.get("relevance_score", 0.0) for chunk in chunks) / len(chunks)
                    document.relevance_score = round(avg_relevance, 2)
                self.db.commit()
                
                # MANDATORY WEAVIATE VALIDATION: Verify the document is actually searchable
                try:
                    # Test that the document can be found in Weaviate
                    search_results = await self.vector_service.hybrid_search(
                        query=title,
                        filters={"tenant_id": tenant_id, "document_code": document_code},
                        limit=1,
                        threshold=0.0
                    )
                    
                    if not search_results:
                        # CRITICAL: Document not searchable in Weaviate - fail the upload
                        print(f"[DocumentService] CRITICAL: Document {document_code} not searchable in Weaviate after upload")
                        
                        # Clean up - delete document and file
                        self.db.delete(document)
                        self.db.commit()
                        await self.storage_service.delete_file(file_path)
                        
                        raise HTTPException(
                            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                            detail="Document upload failed: Document not searchable in Weaviate. This is required for agent access."
                        )
                    
                    print(f"[DocumentService] ✅ Weaviate search validation passed for {document_code}")
                    
                except Exception as search_error:
                    print(f"[DocumentService] CRITICAL: Weaviate search validation failed for {document_code}: {search_error}")
                    
                    # Clean up - delete document and file
                    self.db.delete(document)
                    self.db.commit()
                    await self.storage_service.delete_file(file_path)
                    
                    raise HTTPException(
                        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                        detail=f"Document upload failed: Weaviate validation error - {str(search_error)}"
                    )
            else:
                # CRITICAL: No chunk IDs returned - Weaviate sync failed
                print(f"[DocumentService] CRITICAL: No vector chunks created for {document_code}")
                
                # Clean up - delete document and file  
                self.db.delete(document)
                self.db.commit()
                await self.storage_service.delete_file(file_path)
                
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail="Document upload failed: No vector embeddings created. Weaviate sync is mandatory."
                )
            
            return document
            
        except Exception as e:
            # Cleanup on error - delete partially uploaded document and file
            cleanup_errors = []
            
            # Delete from database if document was created
            if 'document' in locals() and document and document.id:
                try:
                    self.db.delete(document)
                    self.db.commit()
                    print(f"[DocumentService] Cleaned up database record for document {document.id}")
                except Exception as db_error:
                    cleanup_errors.append(f"Database cleanup failed: {str(db_error)}")
                    self.db.rollback()
            
            # Delete file from storage if it was uploaded
            if 'file_path' in locals() and file_path:
                try:
                    await self.storage_service.delete_file(file_path)
                    print(f"[DocumentService] Cleaned up file from storage: {file_path}")
                except Exception as storage_error:
                    cleanup_errors.append(f"Storage cleanup failed: {str(storage_error)}")
            
            # Delete from vector database if chunks were added
            if 'document_code' in locals() and document_code:
                try:
                    await self.vector_service.delete_document_chunks(document_code, tenant_id)
                    print(f"[DocumentService] Cleaned up vector chunks for document {document_code}")
                except Exception as vector_error:
                    cleanup_errors.append(f"Vector cleanup failed: {str(vector_error)}")
            
            # Build error message
            error_message = f"Failed to process document: {str(e)}"
            if cleanup_errors:
                error_message += f" (Cleanup errors: {'; '.join(cleanup_errors)})"
            
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=error_message
            )
    
    async def _extract_text_content(self, file: UploadFile, file_extension: str) -> str:
        """
        Extract text content from uploaded file
        """
        content = ""
        
        try:
            # Always reset file pointer to beginning before reading
            await file.seek(0)
            
            if file_extension == '.pdf':
                # Extract from PDF
                pdf_reader = pypdf.PdfReader(file.file)
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n"
            
            elif file_extension in ['.docx']:
                # Extract from DOCX
                doc = docx.Document(file.file)
                for paragraph in doc.paragraphs:
                    content += paragraph.text + "\n"
            
            elif file_extension == '.txt':
                # Read text file - use proper file content reading
                file_content = await file.read()
                if isinstance(file_content, bytes):
                    content = file_content.decode('utf-8', errors='ignore')
                else:
                    content = str(file_content)
            
            else:
                # Handle other text-based files
                try:
                    file_content = await file.read()
                    if isinstance(file_content, bytes):
                        content = file_content.decode('utf-8', errors='ignore')
                    else:
                        content = str(file_content)
                except UnicodeDecodeError:
                    # Fallback for binary files that can't be decoded
                    content = f"Binary file: {file.filename}"
            
            # Reset file pointer for any subsequent operations
            await file.seek(0)
            
        except Exception as e:
            print(f"[DocumentService] Text extraction error for {file.filename}: {e}")
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Failed to extract text from file: {str(e)}"
            )
        
        return content.strip()
    
    async def _create_semantic_chunks(
        self, 
        content: str, 
        document_type: str, 
        document_code: str
    ) -> List[Dict[str, Any]]:
        """
        Create semantic chunks based on document type with parallel processing
        """
        chunks = []
        
        if document_type == "normativa":
            # For legal documents: split by articles/sections
            chunks = self._legal_document_chunker(content, document_code)
        elif document_type == "istruzione_operativa":
            # For procedures: split by operational steps
            chunks = self._procedure_chunker(content, document_code)
        else:
            # Default chunking
            chunks = self._default_chunker(content, document_code)
        
        # Process chunks in parallel for keyword extraction and scoring
        import asyncio
        
        # Add relevance score to chunks without keywords
        for chunk in chunks:
            chunk["relevance_score"] = self._calculate_relevance_score(chunk["content"])
        
        return chunks
    
    def _legal_document_chunker(self, content: str, document_code: str) -> List[Dict[str, Any]]:
        """
        Chunk legal documents by articles and sections with optimized sizes
        """
        chunks = []
        max_chunk_words = 1500  # Increased for better context
        
        # Split by articles
        article_pattern = r'(Art\.\s*\d+|Articolo\s+\d+)'
        sections = re.split(article_pattern, content, flags=re.IGNORECASE)
        
        current_section = ""
        accumulated_content = ""
        accumulated_sections = []
        
        for i, section in enumerate(sections):
            if re.match(article_pattern, section, re.IGNORECASE):
                # If we have accumulated content, save it as a chunk
                if accumulated_content and len(accumulated_content.split()) > 100:
                    chunks.append({
                        "content": accumulated_content.strip(),
                        "section_title": ", ".join(accumulated_sections),
                        "chunk_index": len(chunks),
                        "word_count": len(accumulated_content.split())
                    })
                    accumulated_content = ""
                    accumulated_sections = []
                
                current_section = section.strip()
            elif section.strip() and current_section:
                new_content = f"{current_section}\n{section.strip()}"
                
                # Check if adding this section would exceed max size
                if accumulated_content and (len(accumulated_content.split()) + len(new_content.split()) > max_chunk_words):
                    # Save current accumulation as chunk
                    chunks.append({
                        "content": accumulated_content.strip(),
                        "section_title": ", ".join(accumulated_sections),
                        "chunk_index": len(chunks),
                        "word_count": len(accumulated_content.split())
                    })
                    accumulated_content = new_content
                    accumulated_sections = [current_section]
                else:
                    # Accumulate content
                    if accumulated_content:
                        accumulated_content += "\n\n" + new_content
                    else:
                        accumulated_content = new_content
                    accumulated_sections.append(current_section)
        
        # Don't forget the last chunk
        if accumulated_content and len(accumulated_content.split()) > 100:
            chunks.append({
                "content": accumulated_content.strip(),
                "section_title": ", ".join(accumulated_sections),
                "chunk_index": len(chunks),
                "word_count": len(accumulated_content.split())
            })
        
        # If no articles found, use optimized paragraph chunking
        if not chunks:
            return self._default_chunker(content, document_code)
        
        return chunks
    
    def _procedure_chunker(self, content: str, document_code: str) -> List[Dict[str, Any]]:
        """
        Chunk operational procedures by steps
        """
        chunks = []
        
        # Split by procedure steps
        step_pattern = r'(Step\s*\d+|Fase\s*\d+|Procedura\s*\d+|\d+\.\s)'
        sections = re.split(step_pattern, content, flags=re.IGNORECASE)
        
        current_step = ""
        for i, section in enumerate(sections):
            if re.match(step_pattern, section, re.IGNORECASE):
                current_step = section.strip()
            elif section.strip() and current_step:
                chunk_content = f"{current_step}\n{section.strip()}"
                if len(chunk_content) > 50:
                    chunks.append({
                        "content": chunk_content,
                        "section_title": current_step,
                        "chunk_index": len(chunks),
                        "word_count": len(chunk_content.split())
                    })
        
        # If no steps found, use default chunking
        if not chunks:
            return self._default_chunker(content, document_code)
        
        return chunks
    
    def _default_chunker(self, content: str, document_code: str) -> List[Dict[str, Any]]:
        """
        Default chunking strategy with optimized sizes
        """
        chunks = []
        max_chunk_size = 1500  # Increased from 1000
        overlap = 200  # Increased from 100 for better context
        
        words = content.split()
        
        for i in range(0, len(words), max_chunk_size - overlap):
            chunk_words = words[i:i + max_chunk_size]
            chunk_content = ' '.join(chunk_words)
            
            if len(chunk_content.strip()) > 100:  # Increased minimum size
                chunks.append({
                    "content": chunk_content,
                    "section_title": f"Sezione {len(chunks) + 1}",
                    "chunk_index": len(chunks),
                    "word_count": len(chunk_words)
                })
        
        return chunks
    
    
    def _calculate_relevance_score(self, text: str) -> float:
        """
        Calculate relevance score - HSE documents start high with differentiation
        """
        if not text:
            return 0.6  # Default good score for HSE context
            
        text_lower = text.lower()
        
        # Start with high base score for HSE documents
        score = 0.7
        
        # Premium terms that boost to excellent scores
        premium_terms = {
            "permesso di lavoro": 0.15, "work permit": 0.15,
            "dpi": 0.12, "ppe": 0.12, "dispositivi protezione": 0.12,
            "procedura sicurezza": 0.12, "safety procedure": 0.12,
            "valutazione rischi": 0.12, "risk assessment": 0.12,
            "lavori a caldo": 0.1, "hot work": 0.1,
            "spazio confinato": 0.1, "confined space": 0.1,
            "lavori elettrici": 0.1, "electrical work": 0.1,
            "cei 11-27": 0.15, "dlgs 81": 0.15, "d.lgs 81": 0.15
        }
        
        # Quality terms that add good value
        quality_terms = {
            "rischio": 0.08, "risk": 0.08, "hazard": 0.08,
            "emergenza": 0.08, "emergency": 0.08,
            "protezione": 0.06, "protection": 0.06,
            "prevenzione": 0.06, "prevention": 0.06,
            "sicurezza": 0.05, "safety": 0.05,
            "controllo": 0.05, "control": 0.05,
            "normativa": 0.07, "standard": 0.07,
            "istruzione": 0.06, "procedure": 0.06
        }
        
        # Count premium terms with higher impact
        for term, weight in premium_terms.items():
            if term in text_lower:
                count = text_lower.count(term)
                score += weight * min(count, 3)  # Cap at 3 occurrences
        
        # Count quality terms
        for term, weight in quality_terms.items():
            if term in text_lower:
                count = text_lower.count(term)
                score += weight * min(count, 2)  # Cap at 2 occurrences
        
        # Bonus for comprehensive HSE content
        term_categories = 0
        if any(term in text_lower for term in ["dpi", "ppe", "protezione"]):
            term_categories += 1
        if any(term in text_lower for term in ["rischio", "risk", "valutazione"]):
            term_categories += 1
        if any(term in text_lower for term in ["procedura", "istruzione", "normativa"]):
            term_categories += 1
        if any(term in text_lower for term in ["lavori", "work", "permesso"]):
            term_categories += 1
            
        if term_categories >= 3:
            score += 0.1  # Comprehensive bonus
        elif term_categories >= 2:
            score += 0.05  # Good coverage bonus
        
        # Ensure excellent scores for clearly HSE documents
        return min(round(score, 3), 1.0)
    
    async def update_document_metadata(
        self,
        document_id: int,
        tenant_id: int,
        title: Optional[str] = None,
        category: Optional[str] = None,
        subcategory: Optional[str] = None,
        industry_sectors: Optional[List[str]] = None,
        authority: Optional[str] = None,
        scope: Optional[str] = None
    ) -> Document:
        """
        Update document metadata without re-uploading file
        """
        document = self.db.query(Document).filter(
            Document.id == document_id,
            Document.tenant_id == tenant_id,
            Document.is_active == True
        ).first()
        
        if not document:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Document not found"
            )
        
        # Update fields if provided
        if title is not None:
            document.title = title
        if category is not None:
            document.category = category
        if subcategory is not None:
            document.subcategory = subcategory
        if industry_sectors is not None:
            document.industry_sectors = industry_sectors
        if authority is not None:
            document.authority = authority
        if scope is not None:
            document.scope = scope
        
        # Update timestamp
        document.updated_at = datetime.utcnow()
        
        self.db.commit()
        self.db.refresh(document)
        
        # Update vector database metadata if title changed
        if title is not None:
            await self.vector_service.update_document_metadata(
                document.document_code, 
                tenant_id, 
                {"title": title}
            )
        
        return document
    
    def get_documents_by_tenant(
        self,
        tenant_id: int,
        document_type: str = None,
        category: str = None,
        limit: int = 50,
        offset: int = 0
    ) -> List[Document]:
        """
        Get documents for a tenant with filtering
        """
        query = self.db.query(Document).filter(
            Document.tenant_id == tenant_id,
            Document.is_active == True
        )
        
        if document_type:
            query = query.filter(Document.document_type == document_type)
        
        if category:
            query = query.filter(Document.category == category)
        
        return query.offset(offset).limit(limit).all()
    
    def get_document_versions(
        self,
        document_code: str,
        tenant_id: int
    ) -> List[Document]:
        """
        Get all versions of a document
        """
        return self.db.query(Document).filter(
            Document.tenant_id == tenant_id,
            Document.document_code == document_code
        ).order_by(Document.version.desc()).all()
    
    def enrich_search_results_with_keywords(
        self,
        search_results: List[Dict[str, Any]],
        tenant_id: int
    ) -> List[Dict[str, Any]]:
        """
        Enrich vector search results with keywords from PostgreSQL
        """
        if not search_results:
            return search_results
        
        # Get unique document codes from search results
        document_codes = list(set(result.get("document_code") for result in search_results if result.get("document_code")))
        
        if not document_codes:
            return search_results
        
        # Fetch documents with keywords from PostgreSQL
        documents = self.db.query(Document).filter(
            Document.tenant_id == tenant_id,
            Document.document_code.in_(document_codes),
            Document.is_active == True
        ).all()
        
        # Create a mapping of document_code to keywords
        keyword_map = {doc.document_code: doc.keywords for doc in documents}
        
        # Enrich search results with keywords
        for result in search_results:
            doc_code = result.get("document_code")
            if doc_code and doc_code in keyword_map:
                result["keywords"] = keyword_map[doc_code]
            else:
                result["keywords"] = []
        
        return search_results
    
    async def delete_document(self, document_id: int, tenant_id: int) -> bool:
        """
        Delete document and its vector chunks
        """
        document = self.db.query(Document).filter(
            Document.id == document_id,
            Document.tenant_id == tenant_id
        ).first()
        
        if not document:
            return False
        
        try:
            # Delete from vector database
            await self.vector_service.delete_document_chunks(
                document.document_code, tenant_id
            )
            
            # Delete file from storage
            if document.file_path:
                await self.storage_service.delete_file(document.file_path)
            
            # Delete from database
            self.db.delete(document)
            self.db.commit()
            
            return True
            
        except Exception as e:
            print(f"Error deleting document: {e}")
            return False
    
    async def _extract_ai_metadata(
        self, 
        content: str, 
        title: str, 
        document_type: str
    ) -> Dict[str, Any]:
        """
        Extract document metadata using AI
        """
        # Truncate content for AI analysis (first 3000 chars to stay within token limits)
        analysis_content = content[:3000] + "..." if len(content) > 3000 else content
        
        prompt = f"""
        Analizza il seguente documento e estrai i metadati richiesti.

        TITOLO: {title}
        TIPO DOCUMENTO: {document_type}
        
        CONTENUTO:
        {analysis_content}

        Estrai e restituisci SOLO un JSON con questa struttura esatta:
        {{
            "category": "string - categoria principale del documento (es: sicurezza, ambiente, qualità, formazione, procedure)",
            "subcategory": "string o null - sottocategoria specifica se applicabile",
            "authority": "string o null - autorità/ente che ha emesso il documento se identificabile",
            "scope": "string - 'tenant' se specifico per l'azienda, 'global' se generico/normativo",
            "industry_sectors": ["array di stringi - settori industriali applicabili (es: edilizia, manifatturiero, chimico, energia)"],
            "keywords": ["array di massimo 15 parole chiave HSE rilevanti estratte dal documento (es: sicurezza, dpi, rischio, prevenzione, emergenza, formazione, manutenzione, controllo, verifica, ambiente, salute)"]
        }}

        Rispondi SOLO con il JSON, senza altri testi.
        """
        
        try:
            # Configure AI provider
            ai_provider = getattr(settings, 'ai_provider', 'gemini').lower()
            
            if ai_provider == 'gemini' and genai:
                # Use Gemini
                if settings.gemini_api_key:
                    genai.configure(api_key=settings.gemini_api_key)
                    model = genai.GenerativeModel(settings.gemini_model or 'gemini-1.5-flash')
                    
                    response = await asyncio.create_task(
                        asyncio.to_thread(model.generate_content, prompt)
                    )
                    ai_response = response.text.strip()
                else:
                    raise ValueError("GEMINI_API_KEY is required")
                    
            elif openai:
                # Use OpenAI as fallback
                api_key = settings.openai_api_key or settings.gemini_api_key
                if not api_key:
                    raise ValueError("Either OPENAI_API_KEY or GEMINI_API_KEY is required")
                    
                client = openai.OpenAI(api_key=api_key)
                response = client.chat.completions.create(
                    model=settings.openai_model or "gpt-3.5-turbo",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.1
                )
                ai_response = response.choices[0].message.content.strip()
            else:
                # Fallback to default values
                return {
                    "category": "generale",
                    "subcategory": None,
                    "authority": None,
                    "scope": "tenant",
                    "industry_sectors": []
                }
            
            # Parse JSON response
            # Clean response to extract only JSON
            start_idx = ai_response.find('{')
            end_idx = ai_response.rfind('}')
            if start_idx != -1 and end_idx != -1:
                json_str = ai_response[start_idx:end_idx+1]
                metadata = json.loads(json_str)
                return metadata
            else:
                raise ValueError("No valid JSON found in AI response")
                
        except Exception as e:
            print(f"Error extracting AI metadata: {e}")
            # Return default values on error
            return {
                "category": "generale",
                "subcategory": None,
                "authority": None,
                "scope": "tenant",
                "industry_sectors": [],
                "keywords": []
            }